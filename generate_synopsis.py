"""
Generate a Project Synopsis for Handwritten Digit Recognition
in the same format as the reference Synopsis document.

Format: A4, Times New Roman, single-column academic synopsis
with cover page, companion table, literature review table,
and standard sections (Abstract, Introduction, Problem Statement,
Objectives, Methodology, Expected Outcomes, References).
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(SCRIPT_DIR, "Synopsis.docx")
IMAGES_DIR = os.path.join(SCRIPT_DIR, "sample_images")


# ── Utility helpers ──────────────────────────────────────────────────

def add_run(paragraph, text, font_name="Times New Roman", font_size=12,
            bold=False, italic=False, underline=False, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure font applies on all character sets
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(
            f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" '
            f'w:hAnsi="{font_name}" w:cs="{font_name}"/>'
        )
        rPr.insert(0, rFonts)
    return run


def set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, props in [("top", top), ("bottom", bottom),
                        ("left", left), ("right", right)]:
        if props:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{props.get("val","single")}" '
                f'w:sz="{props.get("sz","4")}" w:space="0" '
                f'w:color="{props.get("color","000000")}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number(doc, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    """Add a right-aligned page number paragraph."""
    # We skip explicit page numbering since Word will handle in print
    pass


# ── Synopsis Builder ─────────────────────────────────────────────────

class SynopsisBuilder:
    """Build a project synopsis document matching the reference format."""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._page_num = 0

    def _setup_styles(self):
        """Configure default styles and page setup."""
        # Page setup: A4
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.92)
        section.bottom_margin = Inches(0.71)
        section.left_margin = Inches(0.92)
        section.right_margin = Inches(0.35)

        # Default style
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(14)

    def add_new_section(self):
        """Add a page break via new section."""
        new_section = self.doc.add_section()
        new_section.page_width = Inches(8.27)
        new_section.page_height = Inches(11.69)
        new_section.top_margin = Inches(0.92)
        new_section.bottom_margin = Inches(0.71)
        new_section.left_margin = Inches(0.92)
        new_section.right_margin = Inches(0.35)

    # ── Cover Page ──

    def build_cover_page(self):
        """Build the cover/title page."""
        # Spacing at top
        p = self.doc.add_paragraph()
        set_spacing(p, before=0, after=20)

        # Project Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=10, after=10, line=30)
        add_run(p, "Handwritten Digit Recognition Using\n"
                   "Convolutional Neural Networks",
                font_size=24, bold=True)

        # PROJECT SYNOPSIS
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=20, after=6, line=18)
        add_run(p, "PROJECT SYNOPSIS", font_size=14, bold=True)

        # BACHELOR OF TECHNOLOGY
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=6, after=6, line=18)
        add_run(p, "BACHELOR OF TECHNOLOGY", font_size=14, bold=True)

        # Department
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=4, after=12, line=18)
        add_run(p, "Computer Science and Engineering", font_size=15.5)

        # SUBMITTED BY
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=20, after=6, line=16)
        add_run(p, "SUBMITTED BY", font_size=12)

        # Author name
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=4, after=4, line=18)
        add_run(p, "Abhay Singh", font_size=14, bold=True)

        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=10, after=10, line=16)
        add_run(p, "MARCH 2026", font_size=12)

        # SUPERVISOR
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=10, after=4, line=18)
        add_run(p, "SUPERVISOR", font_size=14)

        # Supervisor name — fill in if known, else placeholder
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=4, after=20, line=18)
        add_run(p, "[Supervisor Name]", font_size=14, bold=True)

        # Department & University
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=20, after=4, line=20)
        add_run(p, "DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\n"
                   "COER UNIVERSITY",
                font_size=15.5, bold=True)

        # City
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=4, after=6, line=20)
        add_run(p, "ROORKEE, UTTARAKHAND", font_size=15.5, bold=True)

    # ── Section Heading ──

    def section_heading(self, title):
        """Add a major section heading like 'ABSTRACT :-'"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=14, after=8, line=20)
        add_run(p, title, font_size=15.5, bold=True)

    def sub_heading(self, title):
        """Add sub-heading like 'Data Collection:'"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=8, after=4, line=16)
        add_run(p, title, font_size=14, bold=True)

    # ── Body Text ──

    def body_text(self, text):
        """Add justified body text paragraph."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=2, after=4, line=16)
        add_run(p, text, font_size=12)

    def body_text_left(self, text):
        """Add left-aligned body text."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=2, after=4, line=16)
        add_run(p, text, font_size=12)

    def bold_bullet(self, text):
        """Add a bold bullet/dash item."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=2, after=2, line=16)
        p.paragraph_format.left_indent = Inches(0.3)
        add_run(p, f"\u2013 {text}", font_size=12, bold=True)

    def numbered_item(self, num, text, bold=False):
        """Add a numbered item."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p, before=2, after=2, line=16)
        p.paragraph_format.left_indent = Inches(0.3)
        add_run(p, f"{num}. {text}", font_size=12, bold=bold)

    def bold_item(self, label, description):
        """Add a bold-label item with normal description."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=2, after=2, line=16)
        p.paragraph_format.left_indent = Inches(0.3)
        add_run(p, f"\u2013 {label}: ", font_size=12, bold=True)
        add_run(p, description, font_size=12)

    # ── Tables ──

    def add_companion_table(self):
        """Add the companion/summary table."""
        self.section_heading("COMPANION TABLE :-")

        data = [
            ["Section", "Details"],
            ["Project Title", "Handwritten Digit Recognition Using Convolutional Neural Networks"],
            ["Author", "Abhay Singh"],
            ["Supervisor", "[Supervisor Name]"],
            ["Institution", "Department of Computer Science, COER University, Roorkee"],
            ["Abstract", "A CNN-based system for recognising handwritten digits (0\u20139) using "
                         "the MNIST dataset, achieving 98.91% test accuracy with real-time "
                         "deployment via Streamlit."],
            ["Introduction", "Covers the significance of handwritten digit recognition in "
                             "postal automation, banking, and document digitisation, and the "
                             "advantage of CNNs over traditional methods."],
            ["Problem Statement", "Build an accurate digit classification system using CNNs "
                                  "that maintains spatial features and deploys in real-time."],
            ["Motivation", "HDR is a gateway problem in deep learning with direct commercial "
                           "applications and serves as a foundation for more complex OCR tasks."],
            ["Objectives", "Design CNN architecture, evaluate performance, analyse components, "
                           "deploy web application, and propose improvements."],
            ["Methodology", "Data preprocessing \u2192 CNN model design \u2192 training with "
                            "regularisation \u2192 evaluation \u2192 Streamlit deployment."],
            ["Expected Outcomes", "High-accuracy digit classifier, real-time web app, "
                                  "comprehensive performance analysis, and deployment-ready model."],
        ]

        table = self.doc.add_table(rows=len(data), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(data):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                set_spacing(p, before=2, after=2, line=14)

                if r_idx == 0:
                    # Header row
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run(p, cell_text, font_size=12, bold=True)
                    set_cell_shading(cell, "D9E2F3")
                    border = {"sz": "6", "color": "000000", "val": "single"}
                    set_cell_borders(cell, top=border, bottom=border,
                                     left=border, right=border)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if c_idx == 0:
                        add_run(p, cell_text, font_size=11, bold=True)
                    else:
                        add_run(p, cell_text, font_size=11)
                    border = {"sz": "4", "color": "000000", "val": "single"}
                    set_cell_borders(cell, top=border, bottom=border,
                                     left=border, right=border)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(5.0)

    def add_literature_review_table(self):
        """Add literature review comparison table."""
        data = [
            ["#", "Author(s) & Title", "Objective", "Key Findings", "Methodology"],
            ["1",
             "Y. LeCun et al., Gradient-Based Learning Applied to Document Recognition (1998)",
             "Introduce CNN for document recognition",
             "LeNet-5 achieved 0.8% error on MNIST, established the CNN paradigm",
             "CNN with convolutional and pooling layers"],
            ["2",
             "D. Cire\u015fan et al., Deep Big Simple Neural Nets for Handwritten Digit Recognition (2010)",
             "Push accuracy limits with deeper CNNs",
             "Achieved 0.35% error rate using deep multi-column architecture",
             "Deep CNN with committee of networks"],
            ["3",
             "S. Ioffe & C. Szegedy, Batch Normalization (2015)",
             "Accelerate deep network training",
             "Batch normalisation enables higher learning rates and acts as regulariser",
             "Normalising layer inputs within mini-batches"],
            ["4",
             "N. Srivastava et al., Dropout: A Simple Way to Prevent Overfitting (2014)",
             "Reduce overfitting in deep networks",
             "Dropout creates implicit ensemble, significantly improving generalisation",
             "Randomly zeroing activations during training"],
            ["5",
             "L. Wan et al., Regularization Using DropConnect (2013)",
             "Generalise dropout to weight-level",
             "Achieved 0.21% error on MNIST, outperforming standard dropout",
             "DropConnect on network weights"],
            ["6",
             "S. Sabour et al., Dynamic Routing Between Capsules (2017)",
             "Capture part-whole relationships",
             "Capsule networks achieved 0.25% error with better spatial awareness",
             "Capsule layers with dynamic routing"],
            ["7",
             "V. Vapnik, The Nature of Statistical Learning Theory (1995)",
             "Apply SVM to pattern classification",
             "SVM with RBF kernel achieved 1.4% error, strong for non-neural baseline",
             "Support Vector Machines"],
        ]

        table = self.doc.add_table(rows=len(data), cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row_data in enumerate(data):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                set_spacing(p, before=1, after=1, line=12)

                if r_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_run(p, cell_text, font_size=9, bold=True)
                    set_cell_shading(cell, "D9E2F3")
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_run(p, cell_text, font_size=9)

                border = {"sz": "4", "color": "000000", "val": "single"}
                set_cell_borders(cell, top=border, bottom=border,
                                 left=border, right=border)

        # Set column widths
        widths = [Inches(0.35), Inches(1.8), Inches(1.4), Inches(1.8), Inches(1.4)]
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = w

    # ── Figure ──

    def add_figure(self, path, caption):
        """Add a centered figure with caption."""
        if not os.path.exists(path):
            self.body_text(f"[Figure: {caption}]")
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=8, after=4)
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=2, after=8, line=14)
        add_run(p, caption, font_size=10, italic=True)

    # ── References ──

    def add_reference(self, text):
        """Add a reference entry."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=1, after=3, line=14)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        add_run(p, text, font_size=12, bold=True)

    # ── Page Break ──

    def page_break(self):
        """Insert a page break."""
        self.doc.add_page_break()

    # ── Save ──

    def save(self):
        """Save the document."""
        self.doc.save(DOCX_PATH)
        file_size = os.path.getsize(DOCX_PATH) / 1024
        print(f"\n{'='*55}")
        print(f"  Synopsis DOCX generated successfully!")
        print(f"  File: {DOCX_PATH}")
        print(f"  Size: {file_size:.1f} KB")
        print(f"{'='*55}\n")


# ═══════════════════════════════════════════════════════════════
# BUILD THE SYNOPSIS
# ═══════════════════════════════════════════════════════════════

def build_synopsis():
    print("[INFO] Building Synopsis DOCX ...")

    s = SynopsisBuilder()

    # ═══════════════════════════════════════════════════════════════
    # PAGE 1: COVER PAGE
    # ═══════════════════════════════════════════════════════════════
    s.build_cover_page()

    # ═══════════════════════════════════════════════════════════════
    # PAGE 2: ABSTRACT
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("ABSTRACT :-")

    s.body_text(
        "Handwritten digit recognition is a fundamental problem in pattern recognition and "
        "computer vision with widespread applications in postal mail sorting, bank cheque "
        "processing, and document digitisation. With the rapid advancement of deep learning "
        "techniques, Convolutional Neural Networks (CNNs) have emerged as the most effective "
        "approach for image classification tasks, including digit recognition."
    )
    s.body_text(
        "This project presents a CNN-based approach for recognising handwritten digits (0\u20139) "
        "using the MNIST benchmark dataset. The proposed model architecture employs two "
        "convolutional layers with ReLU activation, max-pooling for spatial downsampling, "
        "batch normalisation for training stability, and dropout regularisation to prevent "
        "overfitting. The model was implemented using TensorFlow/Keras and trained on 60,000 "
        "labelled images."
    )
    s.body_text(
        "The system achieves a test accuracy of 98.91% on the 10,000-image test set, "
        "demonstrating the effectiveness of the CNN architecture for digit classification. "
        "The trained model is deployed through a Streamlit-based web application that features "
        "an interactive drawing canvas, enabling users to draw digits and receive real-time "
        "predictions with confidence scores."
    )
    s.body_text(
        "This synopsis provides an overview of the project\u2019s objectives, methodology, "
        "expected outcomes, and its significance in both academic research and practical "
        "applications such as postal automation, banking, and mobile computing."
    )

    # ═══════════════════════════════════════════════════════════════
    # PAGE 3: COMPANION TABLE
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.add_companion_table()

    # ═══════════════════════════════════════════════════════════════
    # PAGE 4-5: INTRODUCTION
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("INTRODUCTION :-")

    s.body_text(
        "Handwritten digit recognition (HDR) is one of the most extensively studied problems "
        "in the field of computer vision and machine learning. The task involves classifying "
        "images of individual handwritten digits into one of ten classes (0 through 9). Despite "
        "its apparent simplicity, HDR poses significant challenges due to the inherent variability "
        "in human handwriting \u2014 differences in stroke width, slant, size, and personal writing "
        "style create a vast space of possible representations for each digit."
    )
    s.body_text(
        "The advancement of deep learning, particularly Convolutional Neural Networks (CNNs), "
        "has revolutionised the field of image recognition. Unlike traditional machine learning "
        "approaches that depend on handcrafted features, CNNs automatically learn hierarchical "
        "feature representations from raw pixel data. This eliminates the need for manual "
        "feature engineering and enables superior performance on image classification tasks."
    )
    s.body_text(
        "Traditional methods such as Support Vector Machines (SVMs) and k-Nearest Neighbours "
        "(k-NN) flatten a 28\u00d728 image into a 784-dimensional vector, destroying all spatial "
        "locality information. CNNs maintain spatial structure through local connectivity, "
        "shared weights, and hierarchical feature learning, resulting in significantly improved "
        "recognition accuracy with fewer trainable parameters."
    )

    s.sub_heading("BACKGROUND :-")
    s.body_text(
        "The Modified National Institute of Standards and Technology (MNIST) dataset is the "
        "most widely used benchmark for handwritten digit recognition. It contains 70,000 "
        "greyscale images (60,000 for training, 10,000 for testing) of handwritten digits, "
        "each sized 28\u00d728 pixels. The dataset was created by LeCun et al. in 1998 and has "
        "served as the standard benchmark for evaluating digit recognition algorithms."
    )
    s.body_text(
        "Early neural network approaches, particularly LeCun\u2019s LeNet-5 (1998), established "
        "the CNN paradigm with alternating convolutional and pooling layers followed by fully "
        "connected classification layers, achieving a 0.8% error rate. Modern deep learning "
        "architectures have pushed MNIST accuracy to near-perfect levels, with ensemble methods "
        "achieving error rates as low as 0.17%."
    )

    # Add MNIST sample images
    s.add_figure(
        os.path.join(IMAGES_DIR, "training_samples.png"),
        "Figure 1: Sample handwritten digits from the MNIST training dataset"
    )

    s.sub_heading("MOTIVATION :-")
    s.body_text(
        "The motivation for this project stems from both practical applications and academic "
        "interest in deep learning. Handwritten digit recognition has direct real-world impact "
        "across multiple domains."
    )

    # Business angle
    p = s.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=6, after=2, line=16)
    add_run(p, "\u27A2 ", font_name="Noto Sans Symbols", font_size=12)
    add_run(p, "Practical Applications", font_size=12, bold=True)

    s.body_text_left(
        "From a practical standpoint, handwritten digit recognition enables:"
    )
    s.bold_bullet(
        "Postal automation \u2013 Automated ZIP code reading for high-speed mail sorting, "
        "processing billions of items per year"
    )
    s.bold_bullet(
        "Banking \u2013 Cheque amount recognition and validation, handling millions of "
        "transactions daily"
    )
    s.bold_bullet(
        "Document digitisation \u2013 Converting handwritten forms, surveys, and records "
        "into machine-readable text"
    )
    s.bold_bullet(
        "Mobile computing \u2013 On-device handwriting input recognition for touchscreen keyboards"
    )

    # Technical angle
    p = s.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=6, after=2, line=16)
    add_run(p, "\u27A2 ", font_name="Noto Sans Symbols", font_size=12)
    add_run(p, "Technical Perspective", font_size=12, bold=True)

    s.body_text_left(
        "The project also addresses interesting technical challenges:"
    )
    s.bold_bullet(
        "Feature hierarchy: Learning meaningful representations from raw pixels "
        "without manual feature engineering"
    )
    s.bold_bullet(
        "Regularisation: Preventing overfitting with limited training data using "
        "dropout and batch normalisation"
    )
    s.bold_bullet(
        "Real-time deployment: Building an end-to-end pipeline from model training "
        "to interactive web application"
    )
    s.bold_bullet(
        "Domain adaptation: Bridging the gap between clean MNIST images and "
        "user-drawn canvas input"
    )

    s.body_text(
        "In summary, this project has the dual goal of creating practical value (via "
        "accurate digit recognition in real-world applications) and advancing technical "
        "understanding (via detailed analysis of CNN architectures and regularisation "
        "techniques for image classification)."
    )

    # ═══════════════════════════════════════════════════════════════
    # LITERATURE REVIEW
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("LITERATURE REVIEW :-")

    s.body_text(
        "A comprehensive literature review was conducted to understand the evolution of "
        "handwritten digit recognition techniques, from classical statistical methods to "
        "modern deep learning approaches. The following table summarises key research "
        "contributions in this domain."
    )

    s.add_literature_review_table()

    s.body_text(
        "The literature reveals a clear trend: convolutional architectures consistently "
        "outperform traditional methods for image classification. Key innovations including "
        "batch normalisation, dropout regularisation, and ensemble methods have progressively "
        "reduced error rates on MNIST. However, most existing works focus solely on accuracy "
        "maximisation without addressing practical deployment \u2014 a gap this project aims to fill."
    )

    # ═══════════════════════════════════════════════════════════════
    # PROBLEM STATEMENT
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("PROBLEM STATEMENT :-")

    s.body_text(
        "Given the MNIST dataset of 28\u00d728 greyscale handwritten digit images, the objective "
        "is to build an end-to-end system that:"
    )
    s.numbered_item(1, "Preprocesses raw pixel data (reshaping, normalisation, one-hot encoding) "
                       "for CNN input")
    s.numbered_item(2, "Classifies each digit image into one of 10 classes (0\u20139) with high accuracy")
    s.numbered_item(3, "Employs appropriate regularisation (dropout, batch normalisation) to "
                       "prevent overfitting while maintaining generalisation")
    s.numbered_item(4, "Deploys the trained model as a real-time interactive web application "
                       "where users can draw digits and get instant predictions")
    s.numbered_item(5, "Addresses practical challenges such as:")

    s.bold_bullet("Domain gap: Differences between clean MNIST images and user-drawn "
                  "canvas input (stroke width, resolution, anti-aliasing)")
    s.bold_bullet("Confidence calibration: Providing meaningful probability scores alongside "
                  "predictions for user interpretation")
    s.bold_bullet("Latency: Ensuring sub-5ms inference time for responsive real-time interaction")

    # ═══════════════════════════════════════════════════════════════
    # OBJECTIVES
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("OBJECTIVES :-")

    s.body_text(
        "The primary objectives of this project are as follows:"
    )
    s.numbered_item(1, "To design and implement a CNN architecture optimised for handwritten "
                       "digit classification on the MNIST dataset, using convolutional layers, "
                       "max-pooling, batch normalisation, and dropout regularisation.")
    s.numbered_item(2, "To evaluate the model\u2019s performance comprehensively using metrics "
                       "including accuracy, loss, precision, recall, F1-score, and confusion "
                       "matrix analysis.")
    s.numbered_item(3, "To analyse the role of each architectural component (convolution, "
                       "pooling, dropout, batch normalisation, softmax) in the network\u2019s "
                       "classification performance.")
    s.numbered_item(4, "To deploy the trained model as an interactive real-time web application "
                       "using Streamlit, featuring a drawing canvas for user input and confidence "
                       "score visualisation.")
    s.numbered_item(5, "To propose future improvements including data augmentation, deeper "
                       "architectures, ensemble methods, and mobile deployment via model "
                       "quantisation.")

    # ═══════════════════════════════════════════════════════════════
    # METHODOLOGY
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("METHODOLOGY :-")

    s.sub_heading("Data Collection & Preprocessing:")
    s.body_text(
        "The MNIST dataset is loaded directly via TensorFlow/Keras, providing 60,000 training "
        "and 10,000 test images. Three preprocessing steps are applied:"
    )
    s.bold_bullet("Reshaping: Images reshaped from (N, 28, 28) to (N, 28, 28, 1) to add "
                  "the channel dimension required by convolutional layers.")
    s.bold_bullet("Normalisation: Pixel values scaled from [0, 255] to [0.0, 1.0] for "
                  "numerical stability and faster convergence.")
    s.bold_bullet("One-Hot Encoding: Integer labels converted to binary vectors of length "
                  "10 for categorical cross-entropy loss computation.")

    s.sub_heading("Model Architecture:")
    s.body_text(
        "The CNN architecture consists of a feature extraction block followed by a "
        "classification head:"
    )
    s.bold_bullet("Feature Extraction: Two Conv2D layers (32 and 64 filters, 3\u00d73 kernels, "
                  "ReLU activation) learn hierarchical features from edges to digit-level "
                  "patterns, followed by MaxPooling2D (2\u00d72) for spatial downsampling and "
                  "Dropout (0.25) for regularisation.")
    s.bold_bullet("Classification Head: Flatten layer converts 2D features to 1D vector "
                  "(9,216 units), Dense layer (128 units, ReLU) performs classification, "
                  "BatchNormalization ensures training stability, Dropout (0.50) prevents "
                  "overfitting, and final Dense layer (10 units, Softmax) outputs class "
                  "probabilities.")
    s.body_text("Total parameters: ~1,199,882. Trainable parameters: ~1,199,370.")

    s.sub_heading("Training & Optimisation:")
    s.body_text(
        "The model is trained using the Adam optimiser (learning rate 0.001) with categorical "
        "cross-entropy loss. A batch size of 128 balances gradient noise and computational "
        "efficiency. A 10% validation split monitors generalisation during training. Early "
        "stopping with patience=3 (monitoring validation loss, restoring best weights) prevents "
        "overfitting. The model converges in approximately 9 epochs (~3 minutes on CPU)."
    )

    # Add training curves image
    s.add_figure(
        os.path.join(IMAGES_DIR, "training_curves.png"),
        "Figure 2: Training and validation accuracy (left) and loss (right) across 9 epochs"
    )

    s.sub_heading("Evaluation:")
    s.body_text(
        "Model performance is evaluated using multiple metrics: overall test accuracy and loss, "
        "per-class precision, recall, and F1-score, confusion matrix analysis to identify "
        "systematic misclassification patterns, and per-class accuracy visualisation. "
        "Computational efficiency is also assessed (inference time, model size, memory footprint)."
    )

    # Add confusion matrix
    s.add_figure(
        os.path.join(IMAGES_DIR, "confusion_matrix.png"),
        "Figure 3: Confusion matrix showing correct and incorrect predictions for each digit class"
    )

    # Add per-class accuracy
    s.add_figure(
        os.path.join(IMAGES_DIR, "per_class_accuracy.png"),
        "Figure 4: Per-digit accuracy. Digit 1 is highest (99.47%), Digit 9 is lowest (97.92%)"
    )

    s.sub_heading("Deployment:")
    s.body_text(
        "The trained model is deployed as a Streamlit web application featuring an HTML5 "
        "drawing canvas. The deployment pipeline processes user input through RGBA-to-greyscale "
        "conversion, Lanczos resizing to 28\u00d728, normalisation, and CNN inference, displaying "
        "predictions with confidence bar charts in real-time."
    )

    # ═══════════════════════════════════════════════════════════════
    # FLOWCHART
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("FLOWCHART :-")

    # Try to add flowchart image if available, otherwise text-based
    flowchart_path = os.path.join(IMAGES_DIR, "flowchart.png")
    if os.path.exists(flowchart_path):
        s.add_figure(flowchart_path, "Figure 1: Project Methodology Flowchart")
    else:
        # Text-based flowchart
        steps = [
            "MNIST Dataset (60,000 train + 10,000 test images)",
            "\u2193",
            "Data Preprocessing (Reshape \u2192 Normalise \u2192 One-Hot Encode)",
            "\u2193",
            "CNN Model Design (Conv2D \u2192 Conv2D \u2192 MaxPool \u2192 Dropout \u2192 Dense \u2192 BatchNorm \u2192 Dropout \u2192 Softmax)",
            "\u2193",
            "Training (Adam Optimiser, Early Stopping, 9 Epochs)",
            "\u2193",
            "Evaluation (Accuracy: 98.91%, Confusion Matrix, Per-Class Report)",
            "\u2193",
            "Deployment (Streamlit Web App with Drawing Canvas)",
            "\u2193",
            "Real-Time Prediction with Confidence Scores",
        ]
        for step in steps:
            p = s.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_spacing(p, before=3, after=3, line=18)
            if step == "\u2193":
                add_run(p, step, font_size=16, bold=True)
            else:
                add_run(p, step, font_size=12)

    # ═══════════════════════════════════════════════════════════════
    # EXPECTED OUTCOMES
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("EXPECTED OUTCOMES :-")

    s.bold_bullet(
        "High Classification Accuracy \u2013 The CNN model achieves 98.91% test accuracy "
        "on the MNIST dataset, with F1-scores exceeding 0.98 across all digit classes, "
        "demonstrating reliable and consistent digit recognition."
    )
    s.bold_bullet(
        "Detailed Performance Analysis \u2013 Comprehensive evaluation including per-class "
        "precision, recall, F1-scores, confusion matrix visualisation, and analysis of "
        "common misclassification patterns (e.g., 9\u21924 confusion due to similar structures)."
    )
    s.bold_bullet(
        "Effective Regularisation \u2013 The combination of Dropout (0.25 and 0.50) and "
        "Batch Normalisation demonstrates successful prevention of overfitting, with "
        "the training-validation accuracy gap remaining below 0.5%."
    )
    s.bold_bullet(
        "Real-Time Web Application \u2013 A fully functional Streamlit-based web application "
        "enabling users to draw digits on an interactive canvas and receive instant "
        "predictions with confidence score visualisations."
    )
    s.bold_bullet(
        "Fast Inference \u2013 Sub-5ms inference time per image, making the system suitable "
        "for real-time applications on standard CPU hardware."
    )
    s.bold_bullet(
        "Deployment-Ready Model \u2013 A trained model saved in .keras format (~14 MB) with "
        "documented preprocessing pipeline, ready for integration into production systems."
    )
    s.bold_bullet(
        "Reproducibility \u2013 Complete, documented codebase with clear architecture "
        "specification, training configuration, and evaluation methodology that can "
        "be replicated and extended for future research."
    )

    # ═══════════════════════════════════════════════════════════════
    # REFERENCES
    # ═══════════════════════════════════════════════════════════════
    s.page_break()
    s.section_heading("REFERENCES :-")

    refs = [
        "[1] Y. LeCun, L. Bottou, Y. Bengio, and P. Haffner. \u201cGradient-Based Learning "
        "Applied to Document Recognition.\u201d Proceedings of the IEEE, vol. 86, no. 11, "
        "pp. 2278\u20132324, 1998.",

        "[2] Y. LeCun, Y. Bengio, and G. Hinton. \u201cDeep Learning.\u201d Nature, vol. 521, "
        "no. 7553, pp. 436\u2013444, 2015.",

        "[3] I. Goodfellow, Y. Bengio, and A. Courville. Deep Learning. MIT Press, 2016.",

        "[4] D. C. Cire\u015fan, U. Meier, L. M. Gambardella, and J. Schmidhuber. \u201cDeep, "
        "Big, Simple Neural Nets for Handwritten Digit Recognition.\u201d Neural Computation, "
        "vol. 22, no. 12, pp. 3207\u20133220, 2010.",

        "[5] S. Ioffe and C. Szegedy. \u201cBatch Normalization: Accelerating Deep Network "
        "Training by Reducing Internal Covariate Shift.\u201d Proc. ICML, pp. 448\u2013456, 2015.",

        "[6] N. Srivastava, G. Hinton, A. Krizhevsky, I. Sutskever, and R. Salakhutdinov. "
        "\u201cDropout: A Simple Way to Prevent Neural Networks from Overfitting.\u201d JMLR, "
        "vol. 15, pp. 1929\u20131958, 2014.",

        "[7] L. Wan, M. Zeiler, S. Zhang, Y. LeCun, and R. Fergus. \u201cRegularization of "
        "Neural Networks Using DropConnect.\u201d Proc. ICML, pp. 1058\u20131066, 2013.",

        "[8] S. Sabour, N. Frosst, and G. E. Hinton. \u201cDynamic Routing Between Capsules.\u201d "
        "NeurIPS, pp. 3856\u20133866, 2017.",

        "[9] V. Vapnik. The Nature of Statistical Learning Theory. Springer, 1995.",

        "[10] R. Plamondon and S. N. Srihari. \u201cOnline and Off-Line Handwriting Recognition: "
        "A Comprehensive Survey.\u201d IEEE Trans. PAMI, vol. 22, no. 1, pp. 63\u201384, 2000.",

        "[11] A. Krizhevsky, I. Sutskever, and G. E. Hinton. \u201cImageNet Classification with "
        "Deep Convolutional Neural Networks.\u201d NeurIPS, pp. 1097\u20131105, 2012.",

        "[12] D. P. Kingma and J. Ba. \u201cAdam: A Method for Stochastic Optimization.\u201d "
        "Proc. ICLR, 2015.",

        "[13] K. Simonyan and A. Zisserman. \u201cVery Deep Convolutional Networks for "
        "Large-Scale Image Recognition.\u201d Proc. ICLR, 2015.",

        "[14] V. Nair and G. E. Hinton. \u201cRectified Linear Units Improve Restricted "
        "Boltzmann Machines.\u201d Proc. ICML, pp. 807\u2013814, 2010.",

        "[15] D. Scherer, A. M\u00fcller, and S. Behnke. \u201cEvaluation of Pooling Operations "
        "in Convolutional Architectures for Object Recognition.\u201d Proc. ICANN, "
        "pp. 92\u2013101, 2010.",

        "[16] L. K. Hansen and P. Salamon. \u201cNeural Network Ensembles.\u201d IEEE Trans. PAMI, "
        "vol. 12, no. 10, pp. 993\u20131001, 1990.",
    ]
    for ref in refs:
        s.add_reference(ref)

    # ─── Save ──────────────────────────────────────────────────────
    s.save()


if __name__ == "__main__":
    build_synopsis()
