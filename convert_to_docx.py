"""
Convert Research Paper to IEEE Conference Paper format DOCX using python-docx.
Matches the formatting style of IEEE conference papers:
  - Two-column layout (via Word section columns)
  - Times New Roman font throughout
  - Roman numeral section headings (centered, ALL CAPS)
  - Italic subsection headings (A., B., C.)
  - Bold italic abstract with em-dash prefix
  - Compact 8pt references
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(SCRIPT_DIR, "Research_Paper.docx")
IMAGES_DIR = os.path.join(SCRIPT_DIR, "sample_images")


def to_roman(n):
    """Convert integer to Roman numeral."""
    vals = [(1000,'M'),(900,'CM'),(500,'D'),(400,'CD'),
            (100,'C'),(90,'XC'),(50,'L'),(40,'XL'),
            (10,'X'),(9,'IX'),(5,'V'),(4,'IV'),(1,'I')]
    result = ''
    for v, r in vals:
        while n >= v:
            result += r
            n -= v
    return result


def set_cell_shading(cell, color):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set borders on a table cell. Each border arg is a dict: {sz, color, val}."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, props in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if props:
            border = parse_xml(
                f'<w:{edge} {nsdecls("w")} w:val="{props.get("val","single")}" '
                f'w:sz="{props.get("sz","4")}" w:space="0" '
                f'w:color="{props.get("color","000000")}"/>'
            )
            tcBorders.append(border)
    tcPr.append(tcBorders)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)


def add_run(paragraph, text, font_name="Times New Roman", font_size=10,
            bold=False, italic=False, color=None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Ensure Times New Roman works for East Asian text too
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.insert(0, rFonts)
    return run


def set_two_columns(section):
    """Set a section to two-column layout with a gap."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="360"/>')
        sectPr.append(cols)
    else:
        cols.set(qn('w:num'), '2')
        cols.set(qn('w:space'), '360')  # ~0.25 inches gap


class IEEEDocxBuilder:
    """Build an IEEE-formatted Word document."""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._table_counter = 0
        self._fig_counter = 0

    def _setup_page(self):
        """Configure page size And margins for IEEE format."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(0.67)
        section.right_margin = Inches(0.67)

        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(10)
        font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = Pt(12)

    def add_header(self, text):
        """Add IEEE conference header at top of page."""
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(p, text, font_size=8, bold=True, italic=True)
        set_paragraph_spacing(p, before=0, after=0)

    def add_title(self, title):
        """Add centered title in ~24pt Times New Roman."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=6, after=6, line=28)
        add_run(p, title, font_size=24)

    def add_authors(self, authors):
        """Add author block (centered, with italic department)."""
        for author in authors:
            # Name
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=2, after=0, line=14)
            add_run(p, author["name"], font_size=11)

            # Department (italic)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=13)
            add_run(p, author.get("dept", ""), font_size=10, italic=True)

            # University
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=13)
            add_run(p, author.get("univ", ""), font_size=10)

            # City
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=0, line=13)
            add_run(p, author.get("city", ""), font_size=10)

            # Email (italic)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=0, after=6, line=13)
            add_run(p, author.get("email", ""), font_size=10, italic=True)

    def add_abstract(self, text, keywords=None):
        """Add IEEE abstract with bold italic prefix."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=6, after=3, line=11)
        # "Abstract--" prefix in bold italic 9pt
        add_run(p, "Abstract\u2014", font_size=9, bold=True, italic=True)
        # Body in bold 9pt
        add_run(p, text, font_size=9, bold=True)

        if keywords:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_paragraph_spacing(p, before=3, after=6, line=11)
            add_run(p, "Keywords\u2014", font_size=9, bold=True, italic=True)
            add_run(p, keywords, font_size=9, bold=True)

    def start_two_columns(self):
        """Insert a section break and switch to two-column layout."""
        # Add a continuous section break
        new_section = self.doc.add_section()
        new_section.start_type = 1  # Continuous
        new_section.page_width = Inches(8.5)
        new_section.page_height = Inches(11)
        new_section.top_margin = Inches(0.75)
        new_section.bottom_margin = Inches(1.0)
        new_section.left_margin = Inches(0.67)
        new_section.right_margin = Inches(0.67)
        set_two_columns(new_section)

    def section_heading(self, num, title):
        """Add IEEE section heading: centered, Roman numeral, ALL CAPS."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=10, after=4, line=12)
        roman = to_roman(num)
        heading_text = f"{roman}. {title.upper()}"
        add_run(p, heading_text, font_size=10)

    def section_heading_no_number(self, title):
        """Add section heading without number (e.g., REFERENCES)."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=10, after=4, line=12)
        add_run(p, title.upper(), font_size=10)

    def subsection_heading(self, letter, title):
        """Add IEEE subsection: italic, left-aligned, lettered."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=6, after=2, line=12)
        add_run(p, f"{letter}. {title}", font_size=10, italic=True)

    def subsubsection_heading(self, title):
        """Add IEEE subsubsection: italic, indented."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, before=4, after=2, line=12)
        p.paragraph_format.left_indent = Inches(0.15)
        add_run(p, title, font_size=10, italic=True)

    def body_text(self, text):
        """Add justified body text in 10pt Times New Roman."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=1, after=3, line=12)
        p.paragraph_format.first_line_indent = Inches(0.2)
        add_run(p, text, font_size=10)

    def body_text_no_indent(self, text):
        """Add body text without first-line indent."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=1, after=3, line=12)
        add_run(p, text, font_size=10)

    def bullet_item(self, text):
        """Add a bullet point."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=1, after=1, line=12)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, "\u2022 ", font_size=10)
        add_run(p, text, font_size=10)

    def numbered_item(self, num, text):
        """Add a numbered list item."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=1, after=1, line=12)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, f"{num}) ", font_size=10)
        add_run(p, text, font_size=10)

    def code_block(self, text):
        """Add a code block in Courier."""
        for line in text.strip().split("\n"):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, before=0, after=0, line=10)
            p.paragraph_format.left_indent = Inches(0.2)
            add_run(p, line, font_name="Courier New", font_size=7.5)

    def add_table(self, caption, headers, rows):
        """Add IEEE-style table with caption above and minimal borders."""
        self._table_counter += 1
        roman = to_roman(self._table_counter)

        # Table label
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=6, after=1, line=10)
        add_run(p, f"TABLE {roman}", font_size=8)

        # Table caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=3, line=10)
        add_run(p, caption, font_size=8)

        # Create table
        num_cols = len(headers)
        num_rows = len(rows) + 1  # +1 for header
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Style the table - remove all borders first
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
        # Remove default borders
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tblBorders>'
        )
        # Remove existing borders element if any
        existing_borders = tblPr.find(qn('w:tblBorders'))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(borders)

        # Header row
        header_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, before=1, after=1, line=10)
            add_run(p, h, font_size=8, bold=True)
            # Top and bottom border on header
            border_props = {"sz": "6", "color": "000000", "val": "single"}
            set_cell_borders(cell, top=border_props, bottom=border_props)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, cell_text in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
                set_paragraph_spacing(p, before=0, after=0, line=10)
                add_run(p, cell_text, font_size=8)
                # Bottom border only on last row
                if r_idx == len(rows) - 1:
                    border_props = {"sz": "6", "color": "000000", "val": "single"}
                    set_cell_borders(cell, bottom=border_props)

        # Spacing after table
        p = self.doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=3)

    def add_figure(self, path, caption):
        """Add a figure with IEEE-style caption below."""
        self._fig_counter += 1

        if not os.path.exists(path):
            self.body_text_no_indent(f"[Image not found: {os.path.basename(path)}]")
            return

        # Add image centered
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=4, after=2)
        run = p.add_run()
        run.add_picture(path, width=Inches(3.2))

        # Caption
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=1, after=4, line=10)
        add_run(p, f"Fig. {self._fig_counter}. ", font_size=8)
        add_run(p, caption, font_size=8)

    def add_reference(self, ref_text):
        """Add a single reference entry."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_paragraph_spacing(p, before=0, after=2, line=10)
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        add_run(p, ref_text, font_size=8)

    def save(self):
        """Save the document."""
        self.doc.save(DOCX_PATH)
        file_size = os.path.getsize(DOCX_PATH) / 1024
        print(f"\n{'='*55}")
        print(f"  IEEE-format DOCX generated successfully!")
        print(f"  File: {DOCX_PATH}")
        print(f"  Size: {file_size:.1f} KB")
        print(f"{'='*55}\n")


def build_docx():
    """Build the IEEE-formatted Word document."""
    print("[INFO] Building IEEE-format DOCX ...")

    doc = IEEEDocxBuilder()

    # ── Header ──
    doc.add_header("2026 COER University \u2014 Department of Computer Science")

    # ═══════════════════════════════════════════════════════════════
    # TITLE (full width)
    # ═══════════════════════════════════════════════════════════════
    doc.add_title(
        "Handwritten Digit Recognition Using\n"
        "Convolutional Neural Networks:\n"
        "A Deep Learning Approach"
    )

    # ═══════════════════════════════════════════════════════════════
    # AUTHORS
    # ═══════════════════════════════════════════════════════════════
    doc.add_authors([{
        "name": "Abhay Singh",
        "dept": "Department of Computer Science",
        "univ": "COER University",
        "city": "Roorkee, India",
        "email": "abhaychauhan5051a@gmail.com",
    }])

    # ═══════════════════════════════════════════════════════════════
    # ABSTRACT (full width)
    # ═══════════════════════════════════════════════════════════════
    doc.add_abstract(
        "Handwritten digit recognition is a fundamental problem in pattern recognition and computer "
        "vision with widespread applications in postal mail sorting, bank cheque processing, and "
        "document digitisation. This paper presents a Convolutional Neural Network (CNN) based approach "
        "for recognising handwritten digits (0\u20139) using the MNIST benchmark dataset. The proposed model "
        "architecture employs two convolutional layers with ReLU activation, max-pooling for spatial "
        "downsampling, batch normalisation for training stability, and dropout regularisation to prevent "
        "overfitting. The model was implemented using TensorFlow/Keras and trained on 60,000 labelled "
        "images, achieving a test accuracy of 98.91% on the 10,000-image test set. We further demonstrate "
        "a real-time deployment through a Streamlit-based web application featuring an interactive drawing "
        "canvas. This paper provides a comprehensive analysis of the CNN architecture, training dynamics, "
        "performance evaluation, and potential enhancements for production deployment.",
        keywords="Handwritten Digit Recognition, Convolutional Neural Network, Deep Learning, "
                 "MNIST, Image Classification, TensorFlow, Pattern Recognition",
    )

    # ═══════════════════════════════════════════════════════════════
    # BEGIN TWO-COLUMN LAYOUT
    # ═══════════════════════════════════════════════════════════════
    doc.start_two_columns()

    # ── I. INTRODUCTION ──
    doc.section_heading(1, "Introduction")

    doc.subsection_heading("A", "Background")
    doc.body_text(
        "Handwritten digit recognition (HDR) is one of the most extensively studied problems in the "
        "field of computer vision and machine learning. The task involves classifying images of individual "
        "handwritten digits into one of ten classes (0 through 9). Despite its apparent simplicity, HDR "
        "poses significant challenges due to the inherent variability in human handwriting \u2014 differences "
        "in stroke width, slant, size, and personal writing style create a vast space of possible "
        "representations for each digit [1]."
    )
    doc.body_text(
        "The advancement of deep learning, particularly Convolutional Neural Networks (CNNs), has "
        "revolutionised the field of image recognition. CNNs have demonstrated remarkable capabilities "
        "in automatically learning hierarchical feature representations from raw pixel data, eliminating "
        "the need for manual feature engineering that characterised earlier approaches [2]."
    )

    doc.subsection_heading("B", "Problem Statement")
    doc.body_text(
        "Traditional machine learning approaches to handwritten digit recognition \u2014 such as Support Vector "
        "Machines (SVMs) and k-Nearest Neighbours (k-NN) \u2014 depend heavily on handcrafted features and "
        "often fail to capture the spatial relationships inherent in image data. When a 28\u00d728 image is "
        "flattened into a 784-dimensional vector for a fully connected network, all spatial locality "
        "information is lost, leading to suboptimal classification performance and excessive parameter "
        "counts [3]."
    )
    doc.body_text(
        "This paper addresses the problem by leveraging CNNs, which maintain spatial structure through "
        "local connectivity, shared weights, and hierarchical feature learning, resulting in significantly "
        "improved recognition accuracy with fewer trainable parameters."
    )

    doc.subsection_heading("C", "Objectives")
    doc.body_text_no_indent("The objectives of this research are:")
    objectives = [
        "To design and implement a CNN architecture optimised for handwritten digit classification on the MNIST dataset.",
        "To evaluate the model\u2019s performance in terms of accuracy, loss convergence, and generalisation capability.",
        "To analyse the role of each architectural component (convolution, pooling, dropout, batch normalisation).",
        "To deploy the trained model in a real-time web application for practical digit recognition.",
        "To propose improvements and discuss real-world applications.",
    ]
    for i, obj in enumerate(objectives, 1):
        doc.numbered_item(i, obj)

    doc.subsection_heading("D", "Significance")
    doc.body_text_no_indent(
        "Handwritten digit recognition serves as a gateway problem in deep learning education and has "
        "direct real-world impact in:"
    )
    apps = [
        "Postal automation \u2014 Automated ZIP code reading for mail sorting [4]",
        "Banking \u2014 Cheque amount recognition and validation [5]",
        "Document digitisation \u2014 Converting handwritten forms to machine-readable text [6]",
        "Mobile computing \u2014 On-device handwriting input recognition [7]",
    ]
    for a in apps:
        doc.bullet_item(a)

    # ── II. LITERATURE REVIEW ──
    doc.section_heading(2, "Literature Review")

    doc.subsection_heading("A", "Classical Approaches")
    doc.body_text(
        "Early work on digit recognition relied on statistical pattern recognition methods. Vapnik and "
        "colleagues demonstrated the effectiveness of Support Vector Machines (SVMs) on the MNIST dataset, "
        "achieving error rates of approximately 1.4% [8]. k-Nearest Neighbours (k-NN) with various "
        "distance metrics achieved error rates around 3\u20135%, depending on the feature representation "
        "used [9]. Random Forests and Gradient Boosted Decision Trees typically achieved "
        "accuracies between 96\u201397% on MNIST [10]."
    )

    doc.subsection_heading("B", "Neural Network Approaches")
    doc.body_text(
        "The resurgence of neural networks began with LeCun et al.\u2019s seminal work on LeNet-5 (1998), "
        "a pioneering CNN architecture that achieved a 0.8% error rate on MNIST [11]. Multi-Layer "
        "Perceptrons (MLPs) without convolutional layers typically achieved error rates of 1.5\u20133%, "
        "demonstrating the clear advantage of convolutional architectures [12]."
    )

    doc.subsection_heading("C", "Modern Deep Learning Approaches")
    doc.body_text_no_indent(
        "Modern architectures have pushed MNIST accuracy to near-perfect levels. Table I compares "
        "the key methods and their error rates."
    )
    doc.add_table(
        "COMPARISON OF METHODS ON THE MNIST TEST SET",
        ["Method", "Error (%)", "Year", "Ref."],
        [
            ["LeNet-5", "0.80", "1998", "[11]"],
            ["SVM (RBF)", "1.40", "1998", "[8]"],
            ["Deep CNN", "0.35", "2012", "[13]"],
            ["DropConnect", "0.21", "2013", "[14]"],
            ["Batch-Norm CNN", "0.29", "2015", "[15]"],
            ["Ensemble CNNs", "0.17", "2016", "[16]"],
            ["Capsule Net", "0.25", "2017", "[17]"],
        ],
    )

    doc.subsection_heading("D", "Research Gap")
    doc.body_text(
        "While several high-accuracy models exist, most research focuses solely on maximising accuracy "
        "without addressing practical deployment. This paper bridges that gap by (a) providing a "
        "well-documented, reproducible CNN implementation with detailed architectural analysis, and "
        "(b) demonstrating real-time deployment through a web-based interactive application."
    )

    # ── III. DATASET DESCRIPTION ──
    doc.section_heading(3, "Dataset Description")

    doc.subsection_heading("A", "The MNIST Dataset")
    doc.body_text(
        "The Modified National Institute of Standards and Technology (MNIST) dataset [11] is the most "
        "widely used benchmark for handwritten digit recognition. It was created by re-sampling and "
        "normalising digits from the original NIST Special Databases 1 and 3."
    )
    doc.add_table(
        "MNIST DATASET PROPERTIES",
        ["Property", "Value"],
        [
            ["Total images", "70,000"],
            ["Training set", "60,000"],
            ["Test set", "10,000"],
            ["Image dimensions", "28 \u00d7 28 pixels"],
            ["Colour space", "Grayscale"],
            ["Pixel range", "0\u2013255"],
            ["Classes", "10 (digits 0\u20139)"],
        ],
    )

    doc.subsection_heading("B", "Sample Visualisation")
    doc.add_figure(
        os.path.join(IMAGES_DIR, "training_samples.png"),
        "Representative handwritten digits from the MNIST training set.",
    )

    doc.subsection_heading("C", "Data Characteristics")
    doc.body_text(
        "The MNIST dataset is approximately balanced across all 10 classes, with each digit represented "
        "by roughly 5,500\u20137,000 training samples. Images are centre-of-mass aligned and size-normalised "
        "to fit within a 20\u00d720 pixel bounding box. Digits are anti-aliased, producing greyscale pixels "
        "at the edges of strokes."
    )

    # ── IV. METHODOLOGY ──
    doc.section_heading(4, "Methodology")

    doc.subsection_heading("A", "Data Preprocessing")
    doc.body_text_no_indent("Three preprocessing steps were applied:")
    doc.body_text(
        "1) Reshaping: Images reshaped from (N, 28, 28) to (N, 28, 28, 1) to add the channel dimension "
        "required by CNNs. "
        "2) Normalisation: Pixel values scaled from [0, 255] to [0.0, 1.0] for faster convergence [18]. "
        "3) One-Hot Encoding: Integer labels converted to binary vectors of length 10 for categorical "
        "cross-entropy loss."
    )

    doc.subsection_heading("B", "Model Architecture")
    doc.body_text_no_indent(
        "The proposed CNN architecture consists of a feature extraction block followed by a classification head:"
    )
    doc.code_block(
        "Input: 28 x 28 x 1\n"
        "-- Feature Extraction --\n"
        "Conv2D(32, 3x3, ReLU)  -> 26x26x32\n"
        "Conv2D(64, 3x3, ReLU)  -> 24x24x64\n"
        "MaxPool2D(2x2)         -> 12x12x64\n"
        "Dropout(0.25)          -> 12x12x64\n"
        "-- Classification Head --\n"
        "Flatten()              -> 9,216\n"
        "Dense(128, ReLU)       -> 128\n"
        "BatchNorm()            -> 128\n"
        "Dropout(0.50)          -> 128\n"
        "Dense(10, Softmax)     -> 10"
    )
    doc.body_text_no_indent("Total parameters: ~1,199,882. Trainable: ~1,199,370.")

    doc.subsection_heading("C", "Layer-by-Layer Justification")

    doc.subsubsection_heading("1) Convolutional Layers:")
    doc.body_text(
        "The convolution operation applies learnable 3\u00d73 filters across the input, producing feature maps. "
        "The first layer (32 filters) learns low-level features such as edges and corners. The second layer "
        "(64 filters) learns higher-level combinations \u2014 curves, loops, and stroke patterns [19]."
    )

    doc.subsubsection_heading("2) ReLU Activation:")
    doc.body_text(
        "f(x) = max(0, x). Chosen over sigmoid/tanh for computational efficiency, sparse activation, "
        "and gradient flow in deep networks [20]."
    )

    doc.subsubsection_heading("3) Max Pooling:")
    doc.body_text(
        "MaxPooling2D with a 2\u00d72 window performs spatial downsampling, halving both dimensions. This "
        "provides translation invariance and reduces computation by 75% [21]."
    )

    doc.subsubsection_heading("4) Dropout:")
    doc.body_text(
        "Dropout randomly zeroes activations during training. Two layers are used: 0.25 after pooling "
        "for mild regularisation, and 0.50 before output for aggressive regularisation. This implicitly "
        "creates an ensemble of 2^n sub-networks [22]."
    )

    doc.subsubsection_heading("5) Batch Normalisation:")
    doc.body_text(
        "Normalises layer inputs to zero mean and unit variance, enabling higher learning rates and "
        "reducing sensitivity to weight initialisation [15]."
    )

    doc.subsubsection_heading("6) Softmax Output:")
    doc.body_text(
        "Converts raw logits into a probability distribution over 10 classes, with all values summing "
        "to 1.0 for direct interpretation as class probabilities."
    )

    doc.subsection_heading("D", "Training Configuration")
    doc.add_table(
        "TRAINING HYPERPARAMETERS",
        ["Parameter", "Value"],
        [
            ["Optimiser", "Adam [23]"],
            ["Learning rate", "0.001"],
            ["Loss", "Cat. cross-entropy"],
            ["Batch size", "128"],
            ["Max epochs", "15"],
            ["Val. split", "10%"],
            ["Early stopping", "patience=3"],
        ],
    )

    doc.subsection_heading("E", "Implementation Environment")
    doc.add_table(
        "IMPLEMENTATION ENVIRONMENT",
        ["Component", "Specification"],
        [
            ["Language", "Python 3.13"],
            ["Framework", "TensorFlow 2.21 / Keras"],
            ["Hardware", "CPU-based training"],
            ["Training time", "~3 min (9 epochs)"],
            ["Deployment", "Streamlit 1.x"],
        ],
    )

    # ── V. RESULTS AND ANALYSIS ──
    doc.section_heading(5, "Results and Analysis")

    doc.subsection_heading("A", "Training Performance")
    doc.body_text(
        "The model was trained for 9 epochs (early stopping triggered with patience=3, monitoring "
        "validation loss). Training accuracy reached 95.4% in the first epoch and improved monotonically "
        "to 99.3% by epoch 9. The gap between training and validation accuracy remained below 0.5%, "
        "confirming effective regularisation."
    )
    doc.add_figure(
        os.path.join(IMAGES_DIR, "training_curves.png"),
        "Training and validation accuracy (left) and loss (right) across 9 epochs.",
    )

    doc.subsection_heading("B", "Test Set Evaluation")
    doc.add_table(
        "TEST SET EVALUATION RESULTS",
        ["Metric", "Value"],
        [
            ["Test accuracy", "98.91%"],
            ["Test loss", "0.0304"],
            ["Error rate", "1.09%"],
            ["Misclassified", "~109 / 10,000"],
        ],
    )
    doc.body_text(
        "The test accuracy of 98.91% demonstrates strong generalisation. The small gap between "
        "validation (~99.2%) and test accuracy confirms the model generalises well without overfitting."
    )

    doc.subsection_heading("C", "Performance Comparison")
    doc.add_table(
        "PERFORMANCE COMPARISON WITH BASELINE METHODS",
        ["Method", "Acc. (%)", "Params"],
        [
            ["k-NN", "96.9", "\u2014"],
            ["SVM (RBF)", "98.6", "\u2014"],
            ["MLP", "97.8", "~236K"],
            ["Proposed CNN", "98.91", "~1.2M"],
            ["LeNet-5", "99.2", "~60K"],
            ["Deep CNN [13]", "99.65", "~10M"],
        ],
    )

    doc.subsection_heading("D", "Per-Class Classification Report")
    doc.add_table(
        "PER-CLASS CLASSIFICATION REPORT",
        ["Digit", "Prec.", "Recall", "F1", "Supp."],
        [
            ["0", "0.989", "0.993", "0.991", "980"],
            ["1", "0.994", "0.995", "0.994", "1135"],
            ["2", "0.980", "0.994", "0.987", "1032"],
            ["3", "0.989", "0.991", "0.990", "1010"],
            ["4", "0.993", "0.990", "0.991", "982"],
            ["5", "0.976", "0.991", "0.983", "892"],
            ["6", "0.995", "0.980", "0.987", "958"],
            ["7", "0.990", "0.989", "0.990", "1028"],
            ["8", "0.994", "0.988", "0.991", "974"],
            ["9", "0.991", "0.979", "0.985", "1009"],
            ["Avg", "0.989", "0.989", "0.989", "10000"],
        ],
    )
    doc.body_text(
        "Digit 1 achieves the highest recall (99.47%) due to its simple, distinctive stroke. "
        "Digit 9 has the lowest recall (97.92%), frequently confused with digit 4 due to similar "
        "upper-loop structures."
    )

    doc.subsection_heading("E", "Confusion Matrix and Per-Class Accuracy")
    doc.add_figure(
        os.path.join(IMAGES_DIR, "confusion_matrix.png"),
        "Confusion matrix showing correct and incorrect predictions per digit class.",
    )
    doc.add_figure(
        os.path.join(IMAGES_DIR, "per_class_accuracy.png"),
        "Per-digit accuracy. Digit 1 is highest (99.47%), Digit 9 lowest (97.92%).",
    )

    doc.subsection_heading("F", "Analysis of Misclassifications")
    doc.body_text_no_indent("The confusion matrix reveals common misclassification patterns:")
    misclass = [
        "6 \u2192 5: 9 samples, due to similar curved stroke patterns.",
        "9 \u2192 4: 6 samples, upper loop of 9 resembles angular 4.",
        "9 \u2192 5: 6 samples, particularly with open-top writing styles.",
        "0 \u2192 2: 6 samples, due to slanted oval shapes.",
        "4 \u2192 9 and 7 \u2192 2: 5 instances each, structural similarities.",
    ]
    for m in misclass:
        doc.bullet_item(m)
    doc.body_text(
        "These ambiguities often arise from genuinely ambiguous handwriting where even human "
        "annotators may disagree."
    )

    doc.subsection_heading("G", "Computational Efficiency")
    doc.add_table(
        "COMPUTATIONAL PERFORMANCE METRICS",
        ["Metric", "Value"],
        [
            ["Training time", "~3 min (CPU)"],
            ["Inference time", "< 5 ms / image"],
            ["Model size", "~14 MB (.keras)"],
            ["Memory", "~50 MB"],
        ],
    )

    # ── VI. DEPLOYMENT ──
    doc.section_heading(6, "Deployment")

    doc.subsection_heading("A", "Web Application Architecture")
    doc.body_text(
        "The trained CNN model was deployed as a real-time web application using Streamlit. The "
        "application consists of: (1) a frontend with a drawable HTML5 canvas via "
        "streamlit-drawable-canvas, (2) a backend with the Keras model cached for fast inference, "
        "and (3) a processing pipeline: canvas \u2192 PIL image \u2192 grayscale \u2192 resize 28\u00d728 \u2192 "
        "normalise \u2192 CNN prediction \u2192 confidence visualisation."
    )

    doc.subsection_heading("B", "Real-Time Preprocessing")
    doc.body_text(
        "Custom images undergo: RGBA to grayscale conversion, resizing to 28\u00d728 with Lanczos "
        "interpolation, normalisation to [0, 1], and reshaping to (1, 28, 28, 1) tensor."
    )

    # ── VII. DISCUSSION ──
    doc.section_heading(7, "Discussion")

    doc.subsection_heading("A", "Strengths")
    doc.body_text(
        "1) High accuracy (98.91%) with only 9 layers and ~1.2M parameters. "
        "2) Effective regularisation via Dropout + BatchNorm preventing overfitting. "
        "3) Fast training (~3 min CPU) and inference (<5 ms). "
        "4) End-to-end practical deployment through Streamlit."
    )

    doc.subsection_heading("B", "Limitations")
    doc.body_text(
        "1) MNIST-specific training \u2014 performance on unconstrained real-world images would be lower. "
        "2) Single-digit recognition only \u2014 multi-digit reading requires segmentation. "
        "3) Domain gap between canvas-drawn digits and MNIST distribution."
    )

    doc.subsection_heading("C", "Proposed Improvements")
    doc.body_text(
        "1) Data augmentation (rotation \u00b110\u00b0, translation, zoom, shear) for improved robustness [24]. "
        "2) Deeper architecture with additional convolutional blocks for 99.5%+ accuracy. "
        "3) Learning rate scheduling via ReduceLROnPlateau. "
        "4) Ensemble of 3\u20135 models for 0.1\u20130.3% error reduction [16]. "
        "5) Residual connections and Capsule Networks [17] for advanced feature learning."
    )

    # ── VIII. REAL-WORLD APPLICATIONS ──
    doc.section_heading(8, "Real-World Applications")
    doc.add_table(
        "REAL-WORLD APPLICATIONS OF HANDWRITTEN DIGIT RECOGNITION",
        ["Domain", "Application", "Scale"],
        [
            ["Postal", "ZIP code reading", "Billions/yr"],
            ["Banking", "Cheque recognition", "Millions/day"],
            ["Healthcare", "Prescription reading", "Hospital"],
            ["Education", "Exam grading", "Institutional"],
            ["Government", "Tax/census digitisation", "National"],
            ["Transport", "License plate digits", "City-wide"],
            ["Mobile", "Handwriting input", "Billions"],
            ["Security", "CAPTCHA recognition", "Web-scale"],
        ],
    )

    # ── IX. CONCLUSION ──
    doc.section_heading(9, "Conclusion")
    doc.body_text(
        "This paper presented a CNN-based approach for handwritten digit recognition using the MNIST "
        "benchmark dataset. The proposed architecture, comprising two convolutional layers, max-pooling, "
        "batch normalisation, and dropout regularisation, achieved a test accuracy of 98.91% \u2014 "
        "demonstrating that a relatively simple and well-regularised CNN can deliver high accuracy."
    )
    doc.body_text(
        "Key contributions include: (1) a clearly documented and reproducible CNN implementation, "
        "(2) comprehensive analysis of training dynamics and regularisation, (3) practical deployment "
        "in a real-time web application, and (4) discussion of limitations and improvement proposals."
    )
    doc.body_text(
        "Future work will focus on multi-digit recognition, training on more challenging datasets "
        "(EMNIST, SVHN), and exploring lightweight architectures for mobile deployment through model "
        "quantisation and knowledge distillation."
    )

    # ── REFERENCES ──
    doc.section_heading_no_number("References")

    refs = [
        '[1] R. Plamondon and S. N. Srihari, \u201cOnline and off-line handwriting recognition: a comprehensive survey,\u201d IEEE Trans. PAMI, vol. 22, no. 1, pp. 63\u201384, 2000.',
        '[2] Y. LeCun, Y. Bengio, and G. Hinton, \u201cDeep learning,\u201d Nature, vol. 521, pp. 436\u2013444, 2015.',
        '[3] I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning. MIT Press, 2016.',
        '[4] S. Impedovo et al., \u201cOptical character recognition \u2014 a survey,\u201d IJPRAI, vol. 5, pp. 1\u201324, 1991.',
        '[5] G. Dimauro et al., \u201cAutomatic bankcheck processing,\u201d IJPRAI, vol. 11, pp. 467\u2013504, 1997.',
        '[6] V. Mitra and C. J. Acharya, \u201cGesture recognition: A survey,\u201d IEEE Trans. SMC, vol. 37, pp. 311\u2013324, 2007.',
        '[7] A. Graves et al., \u201cA novel connectionist system for unconstrained handwriting recognition,\u201d IEEE Trans. PAMI, vol. 31, pp. 855\u2013868, 2009.',
        '[8] V. Vapnik, The Nature of Statistical Learning Theory. Springer, 1995.',
        '[9] T. M. Cover and P. E. Hart, \u201cNearest neighbor pattern classification,\u201d IEEE Trans. IT, vol. 13, pp. 21\u201327, 1967.',
        '[10] L. Breiman, \u201cRandom forests,\u201d Machine Learning, vol. 45, pp. 5\u201332, 2001.',
        '[11] Y. LeCun et al., \u201cGradient-based learning applied to document recognition,\u201d Proc. IEEE, vol. 86, pp. 2278\u20132324, 1998.',
        '[12] K. Hornik et al., \u201cMultilayer feedforward networks are universal approximators,\u201d Neural Networks, vol. 2, pp. 359\u2013366, 1989.',
        '[13] D. C. Cire\u015fan et al., \u201cDeep, big, simple neural nets for handwritten digit recognition,\u201d Neural Computation, vol. 22, pp. 3207\u20133220, 2010.',
        '[14] L. Wan et al., \u201cRegularization of neural networks using DropConnect,\u201d Proc. ICML, pp. 1058\u20131066, 2013.',
        '[15] S. Ioffe and C. Szegedy, \u201cBatch normalization,\u201d Proc. ICML, pp. 448\u2013456, 2015.',
        '[16] L. K. Hansen and P. Salamon, \u201cNeural network ensembles,\u201d IEEE Trans. PAMI, vol. 12, pp. 993\u20131001, 1990.',
        '[17] S. Sabour et al., \u201cDynamic routing between capsules,\u201d NeurIPS, pp. 3856\u20133866, 2017.',
        '[18] D. P. Kingma and J. Ba, \u201cAdam: A method for stochastic optimization,\u201d Proc. ICLR, 2015.',
        '[19] K. Simonyan and A. Zisserman, \u201cVery deep convolutional networks,\u201d Proc. ICLR, 2015.',
        '[20] V. Nair and G. E. Hinton, \u201cRectified linear units improve restricted Boltzmann machines,\u201d Proc. ICML, pp. 807\u2013814, 2010.',
        '[21] D. Scherer et al., \u201cEvaluation of pooling operations in convolutional architectures,\u201d Proc. ICANN, pp. 92\u2013101, 2010.',
        '[22] N. Srivastava et al., \u201cDropout: a simple way to prevent neural networks from overfitting,\u201d JMLR, vol. 15, pp. 1929\u20131958, 2014.',
        '[23] D. P. Kingma and J. Ba, \u201cAdam: A method for stochastic optimization,\u201d arXiv:1412.6980, 2014.',
        '[24] A. Krizhevsky et al., \u201cImageNet classification with deep convolutional neural networks,\u201d NeurIPS, pp. 1097\u20131105, 2012.',
    ]
    for ref in refs:
        doc.add_reference(ref)

    # ─── Save ──────────────────────────────────────────────────────
    doc.save()


if __name__ == "__main__":
    build_docx()
